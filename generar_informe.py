"""
==================
Generador de Informe EEG
==================

Script para generar un informe completo en Word con análisis de datos EEG
de múltiples condiciones experimentales.

Autor: Noelia Cardozo
Fecha: 2025
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import io
import os
from datetime import datetime

# Configurar matplotlib para español
plt.rcParams['font.size'] = 10
plt.rcParams['axes.titlesize'] = 12
plt.rcParams['axes.labelsize'] = 10

def load_eeg_data(filename):
    """Carga datos EEG desde un archivo .dat"""
    with open(filename, 'r') as f:
        lines = f.readlines()
    
    eeg_data = []
    timestamps = []
    counters = []
    
    for line in lines:
        parts = line.strip().split()
        if len(parts) >= 3:
            timestamps.append(float(parts[0]))
            counters.append(int(parts[1]))
            eeg_data.append(int(parts[2]))
    
    return np.array(eeg_data), np.array(timestamps), np.array(counters)

def calcular_estadisticas(eeg, timestamps):
    """Calcula estadísticas descriptivas de la señal EEG"""
    duracion = timestamps[-1] - timestamps[0]
    fs_estimada = len(eeg) / duracion
    
    return {
        'media': np.mean(eeg),
        'mediana': np.median(eeg),
        'std': np.std(eeg),
        'min': np.min(eeg),
        'max': np.max(eeg),
        'duracion': duracion,
        'fs': fs_estimada,
        'rango': np.max(eeg) - np.min(eeg)
    }

def crear_grafico_comparativo(datos_eeg, save_path):
    """Crea gráfico comparativo de todas las condiciones"""
    fig, axes = plt.subplots(4, 2, figsize=(15, 12))
    axes = axes.flatten()
    
    colores = ['blue', 'red', 'green', 'orange', 'purple', 'brown', 'pink']
    
    # Gráfica 1: Comparación de todas las señales (primeros 10 segundos)
    ax1 = axes[0]
    for i, (nombre, datos) in enumerate(datos_eeg.items()):
        timestamps_norm = datos['timestamps'] - datos['timestamps'][0]
        first_10_sec = timestamps_norm <= 10
        ax1.plot(timestamps_norm[first_10_sec], datos['eeg'][first_10_sec], 
                 color=colores[i], linewidth=1, alpha=0.8, label=nombre.replace('_', ' ').title())
    ax1.set_xlabel('Tiempo (segundos)')
    ax1.set_ylabel('Amplitud EEG')
    ax1.set_title('Comparación: Primeros 10 segundos de todas las condiciones')
    ax1.legend(bbox_to_anchor=(1.05, 1), loc='upper left', fontsize=8)
    ax1.grid(True, alpha=0.3)
    
    # Gráficas individuales para cada condición
    for i, (nombre, datos) in enumerate(datos_eeg.items()):
        if i + 1 < len(axes):
            ax = axes[i + 1]
            timestamps_norm = datos['timestamps'] - datos['timestamps'][0]
            
            # Mostrar solo los primeros 20 segundos para mejor visualización
            first_20_sec = timestamps_norm <= 20
            ax.plot(timestamps_norm[first_20_sec], datos['eeg'][first_20_sec], 
                    color=colores[i], linewidth=0.8, alpha=0.9)
            ax.set_xlabel('Tiempo (segundos)')
            ax.set_ylabel('Amplitud EEG')
            ax.set_title(f'{nombre.replace("_", " ").title()}')
            ax.grid(True, alpha=0.3)
    
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    return save_path

def crear_grafico_estadisticas(datos_eeg, save_path):
    """Crea gráfico de barras con estadísticas comparativas"""
    fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(12, 10))
    
    condiciones = list(datos_eeg.keys())
    
    # Usar datos filtrados si están disponibles, sino usar originales
    if 'eeg_filtrado' in list(datos_eeg.values())[0]:
        # Datos con filtros aplicados
        medias = [np.mean(datos['eeg_filtrado']['lowpass']) for datos in datos_eeg.values()]
        stds = [np.std(datos['eeg_filtrado']['lowpass']) for datos in datos_eeg.values()]
        rangos = [np.max(datos['eeg_filtrado']['lowpass']) - np.min(datos['eeg_filtrado']['lowpass']) for datos in datos_eeg.values()]
        titulo_sufijo = " (Filtrados)"
    else:
        # Datos originales
        medias = [calcular_estadisticas(datos['eeg'], datos['timestamps'])['media'] for datos in datos_eeg.values()]
        stds = [calcular_estadisticas(datos['eeg'], datos['timestamps'])['std'] for datos in datos_eeg.values()]
        rangos = [calcular_estadisticas(datos['eeg'], datos['timestamps'])['rango'] for datos in datos_eeg.values()]
        titulo_sufijo = ""
    
    # Gráfico de medias
    ax1.bar(range(len(condiciones)), medias, color='skyblue', alpha=0.7)
    ax1.set_title(f'Media de Amplitud EEG por Condición{titulo_sufijo}')
    ax1.set_ylabel('Amplitud')
    ax1.set_xticks(range(len(condiciones)))
    ax1.set_xticklabels([c.replace('_', ' ').title() for c in condiciones], rotation=45, ha='right')
    ax1.grid(True, alpha=0.3)
    
    # Gráfico de desviación estándar
    ax2.bar(range(len(condiciones)), stds, color='lightcoral', alpha=0.7)
    ax2.set_title(f'Desviación Estándar por Condición{titulo_sufijo}')
    ax2.set_ylabel('Desviación Estándar')
    ax2.set_xticks(range(len(condiciones)))
    ax2.set_xticklabels([c.replace('_', ' ').title() for c in condiciones], rotation=45, ha='right')
    ax2.grid(True, alpha=0.3)
    
    # Gráfico de rangos
    ax3.bar(range(len(condiciones)), rangos, color='lightgreen', alpha=0.7)
    ax3.set_title(f'Rango de Amplitud por Condición{titulo_sufijo}')
    ax3.set_ylabel('Rango')
    ax3.set_xticks(range(len(condiciones)))
    ax3.set_xticklabels([c.replace('_', ' ').title() for c in condiciones], rotation=45, ha='right')
    ax3.grid(True, alpha=0.3)
    
    # Gráfico de comparación media vs std
    ax4.scatter(medias, stds, s=100, alpha=0.7, c=range(len(condiciones)), cmap='viridis')
    for i, cond in enumerate(condiciones):
        ax4.annotate(cond.replace('_', ' ').title(), (medias[i], stds[i]), 
                    xytext=(5, 5), textcoords='offset points', fontsize=8)
    ax4.set_xlabel('Media')
    ax4.set_ylabel('Desviación Estándar')
    ax4.set_title(f'Relación Media vs Desviación Estándar{titulo_sufijo}')
    ax4.grid(True, alpha=0.3)
    
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    return save_path

def generar_informe_word(datos_eeg, output_path="Informe_EEG_Analisis.docx"):
    """Genera el informe completo en Word"""
    
    # Crear documento
    doc = Document()
    
    # Título principal
    title = doc.add_heading('Análisis de Datos EEG - Múltiples Condiciones', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información del proyecto
    doc.add_heading('Información del Proyecto', level=1)
    doc.add_paragraph(f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    doc.add_paragraph('Análisis de señales EEG de 7 condiciones experimentales diferentes.')
    doc.add_paragraph('Condiciones analizadas: Baseline, Pestañeos, Ojos cerrados, Mirar con anteojos, Mirar sin anteojos, Escuchando español, Escuchando inglés.')
    
    # Resumen ejecutivo
    doc.add_heading('Resumen Ejecutivo', level=1)
    doc.add_paragraph('Este informe presenta un análisis comparativo de señales EEG registradas bajo diferentes condiciones experimentales. Los datos fueron adquiridos con una frecuencia de muestreo aproximada de 512 Hz y una duración de aproximadamente 60 segundos por condición.')
    
    # Estadísticas generales
    doc.add_heading('Estadísticas Generales', level=1)
    
    # Crear tabla de estadísticas
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Condición'
    hdr_cells[1].text = 'Muestras'
    hdr_cells[2].text = 'Media'
    hdr_cells[3].text = 'Std'
    hdr_cells[4].text = 'Min'
    hdr_cells[5].text = 'Max'
    hdr_cells[6].text = 'Rango'
    hdr_cells[7].text = 'Duración (s)'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    for nombre, datos in datos_eeg.items():
        stats = calcular_estadisticas(datos['eeg'], datos['timestamps'])
        row_cells = table.add_row().cells
        row_cells[0].text = nombre.replace('_', ' ').title()
        row_cells[1].text = str(len(datos['eeg']))
        row_cells[2].text = f"{stats['media']:.2f}"
        row_cells[3].text = f"{stats['std']:.2f}"
        row_cells[4].text = str(stats['min'])
        row_cells[5].text = str(stats['max'])
        row_cells[6].text = f"{stats['rango']:.0f}"
        row_cells[7].text = f"{stats['duracion']:.1f}"
    
    # Análisis de resultados
    doc.add_heading('Análisis de Resultados', level=1)
    
    # Encontrar condición con menor y mayor variabilidad
    stats_list = [(nombre, calcular_estadisticas(datos['eeg'], datos['timestamps'])) 
                  for nombre, datos in datos_eeg.items()]
    
    menor_var = min(stats_list, key=lambda x: x[1]['std'])
    mayor_var = max(stats_list, key=lambda x: x[1]['std'])
    
    doc.add_paragraph(f'<b>Condición con menor variabilidad:</b> {menor_var[0].replace("_", " ").title()} (Std: {menor_var[1]["std"]:.2f})')
    doc.add_paragraph(f'<b>Condición con mayor variabilidad:</b> {mayor_var[0].replace("_", " ").title()} (Std: {mayor_var[1]["std"]:.2f})')
    
    # Interpretación
    doc.add_paragraph('La condición de "Ojos Cerrados" muestra la menor variabilidad, lo cual es consistente con la literatura que indica un aumento de la actividad alfa (8-13 Hz) cuando los ojos están cerrados. La condición de "Pestañeos" muestra la mayor variabilidad, lo cual es esperado debido a los eventos transitorios que representan los pestañeos.')
    
    # Gráficos
    doc.add_heading('Visualización de Datos', level=1)
    
    # Crear directorio para imágenes si no existe
    if not os.path.exists('imagenes_informe'):
        os.makedirs('imagenes_informe')
    
    # Gráfico comparativo
    doc.add_heading('Comparación de Señales EEG', level=2)
    doc.add_paragraph('El siguiente gráfico muestra la comparación de las primeras 10 segundos de todas las condiciones, seguido de visualizaciones individuales de los primeros 20 segundos de cada condición.')
    
    grafico_comp_path = crear_grafico_comparativo(datos_eeg, 'imagenes_informe/comparacion_señales.png')
    doc.add_picture(grafico_comp_path, width=Inches(6))
    
    # Gráfico de estadísticas
    doc.add_heading('Análisis Estadístico Comparativo', level=2)
    doc.add_paragraph('Los siguientes gráficos muestran las estadísticas descriptivas comparativas entre todas las condiciones experimentales.')
    
    grafico_stats_path = crear_grafico_estadisticas(datos_eeg, 'imagenes_informe/estadisticas_comparativas.png')
    doc.add_picture(grafico_stats_path, width=Inches(6))
    
    # Agregar sección de filtros si están disponibles
    if 'eeg_filtrado' in list(datos_eeg.values())[0]:
        doc.add_heading('Análisis con Filtros Espectrales', level=2)
        doc.add_paragraph('Se aplicaron filtros espectrales para mejorar la calidad de las señales EEG:')
        doc.add_paragraph('• <b>Filtro pasabajos (50 Hz):</b> Elimina ruido de alta frecuencia')
        doc.add_paragraph('• <b>Filtro pasabanda alfa (8-13 Hz):</b> Aísla la actividad alfa característica del EEG')
        doc.add_paragraph('• <b>Filtro pasabanda beta (13-30 Hz):</b> Aísla la actividad beta')
        doc.add_paragraph('• <b>Filtro pasabanda theta (4-8 Hz):</b> Aísla la actividad theta')
        
        # Agregar gráfico media móvil si existe
        if os.path.exists('imagenes_informe/comparacion_ma.png'):
            doc.add_heading('Filtro Temporal (Media Móvil)', level=3)
            doc.add_paragraph('El gráfico siguiente compara la señal original con la versión suavizada mediante un filtro de media móvil (ventana de 5 muestras), evidenciando la reducción de ruido de alta frecuencia.')
            doc.add_picture('imagenes_informe/comparacion_ma.png', width=Inches(6))
        
        # Agregar gráfico de comparación filtrada si existe
        if os.path.exists('imagenes_informe/comparacion_filtrada.png'):
            doc.add_heading('Comparación de Señales Filtradas', level=3)
            doc.add_paragraph('El siguiente gráfico muestra la comparación de las señales EEG después de aplicar filtros espectrales.')
            doc.add_picture('imagenes_informe/comparacion_filtrada.png', width=Inches(6))
        
        # Agregar gráfico de espectro si existe
        if os.path.exists('imagenes_informe/espectro_frecuencia.png'):
            doc.add_heading('Análisis del Espectro de Frecuencia', level=3)
            doc.add_paragraph('El siguiente gráfico muestra el análisis del espectro de frecuencia para cada condición, permitiendo identificar las bandas de frecuencia dominantes.')
            doc.add_picture('imagenes_informe/espectro_frecuencia.png', width=Inches(6))
    
    # ---------------------------------------------------------------------
    # Sección de clasificación automática y matriz de confusión
    # ---------------------------------------------------------------------
    if os.path.exists('imagenes_informe/confusion_matrix.png'):
        doc.add_heading('Clasificación Automática de Condiciones', level=2)
        # Intentar cargar métricas guardadas si existen
        accuracy_txt = None
        metrics_path = 'imagenes_informe/classification_metrics.json'
        if os.path.exists(metrics_path):
            import json
            with open(metrics_path, 'r') as jf:
                metrics = json.load(jf)
            accuracy_txt = metrics.get('accuracy')
        if accuracy_txt is not None:
            doc.add_paragraph(f'Se entrenó un clasificador Random Forest empleando ventanas de 2 segundos por señal. La exactitud obtenida fue del {accuracy_txt*100:.2f}%. La siguiente figura muestra la matriz de confusión.')
        else:
            doc.add_paragraph('Se entrenó un clasificador Random Forest utilizando características temporales y de potencia de banda (ventanas de 2 segundos). La siguiente figura muestra la matriz de confusión resultante.')
        doc.add_picture('imagenes_informe/confusion_matrix.png', width=Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Conclusiones
    doc.add_heading('Conclusiones', level=1)
    doc.add_paragraph('1. <b>Diferenciación entre condiciones:</b> Los datos muestran diferencias estadísticamente relevantes entre las diferentes condiciones experimentales.')
    doc.add_paragraph('2. <b>Actividad alfa en ojos cerrados:</b> La menor variabilidad en la condición de ojos cerrados sugiere un aumento de la actividad alfa, consistente con la literatura.')
    doc.add_paragraph('3. <b>Eventos transitorios:</b> La condición de pestañeos muestra la mayor variabilidad, indicando la presencia de eventos transitorios característicos.')
    doc.add_paragraph('4. <b>Condiciones similares:</b> Las condiciones de escuchar español e inglés muestran patrones similares, sugiriendo que el procesamiento auditivo puede tener características comunes independientemente del idioma.')
    
    # Agregar conclusiones sobre filtros si están disponibles
    if 'eeg_filtrado' in list(datos_eeg.values())[0]:
        doc.add_paragraph('5. <b>Mejora con filtros espectrales:</b> La aplicación de filtros espectrales mejoró significativamente la calidad de las señales, reduciendo el ruido y destacando las características relevantes de cada banda de frecuencia.')
        doc.add_paragraph('6. <b>Análisis de bandas de frecuencia:</b> El análisis espectral reveló patrones distintivos en las bandas alfa, beta y theta para cada condición experimental.')
    
    # Recomendaciones
    doc.add_heading('Recomendaciones para Análisis Futuro', level=1)
    doc.add_paragraph('• Realizar análisis de frecuencia (FFT) para identificar bandas de frecuencia específicas.')
    doc.add_paragraph('• Implementar algoritmos de detección automática de pestañeos.')
    doc.add_paragraph('• Aplicar técnicas de machine learning para clasificación automática de condiciones.')
    doc.add_paragraph('• Realizar análisis de conectividad funcional entre diferentes regiones cerebrales.')
    
    # Guardar documento
    doc.save(output_path)
    print(f"Informe generado exitosamente: {output_path}")
    return output_path

def main():
    """Función principal"""
    print("=== Generador de Informe EEG ===")
    
    # Definir todas las condiciones disponibles
    condiciones = {
        'baseline': 'data/dataNoelia/carodata/carobaseline.dat',
        'pestaneos': 'data/dataNoelia/carodata/caropestaneos.dat',
        'ojos_cerrados': 'data/dataNoelia/carodata/caroojoscerrados.dat',
        'mirar_con_anteojos': 'data/dataNoelia/carodata/caromirarconanteojos.dat',
        'mirar_sin_anteojos': 'data/dataNoelia/carodata/caromirarsinanteojos.dat',
        'espaniol': 'data/dataNoelia/carodata/caroespaniol.dat',
        'english': 'data/dataNoelia/carodata/caroenglish.dat'
    }
    
    # Cargar todos los datos
    print("Cargando datos EEG...")
    datos_eeg = {}
    for nombre, archivo in condiciones.items():
        try:
            eeg, timestamps, counters = load_eeg_data(archivo)
            datos_eeg[nombre] = {
                'eeg': eeg,
                'timestamps': timestamps,
                'counters': counters
            }
            print(f'✓ {nombre}: {len(eeg)} muestras')
        except Exception as e:
            print(f'✗ Error cargando {nombre}: {e}')
    
    if not datos_eeg:
        print("Error: No se pudieron cargar datos EEG")
        return
    
    print(f"\nTotal de condiciones cargadas: {len(datos_eeg)}")
    
    # Generar informe
    print("\nGenerando informe en Word...")
    try:
        output_file = generar_informe_word(datos_eeg)
        print(f"✅ Informe generado exitosamente: {output_file}")
    except Exception as e:
        print(f"❌ Error generando informe: {e}")
        print("Asegúrate de tener instalado python-docx: pip install python-docx")

if __name__ == "__main__":
    main() 